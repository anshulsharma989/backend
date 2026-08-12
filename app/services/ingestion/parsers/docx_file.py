"""DOCX parser."""

from pathlib import Path

from docx import Document as DocxDocument

from app.services.ingestion.parsers.base import BaseParser, ParsedSegment, register_parser


@register_parser("docx")
class DocxParser(BaseParser):
    def parse(self, path: Path) -> list[ParsedSegment]:
        doc = DocxDocument(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        text = "\n\n".join(paragraphs)
        return [ParsedSegment(text=text)] if text else []
